#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成面向对象程序设计期中考试实验报告"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('面向对象程序设计期中考试实验')
run.font.size = Pt(22)
run.font.bold = True

doc.add_paragraph()

info_items = [
    ('学生姓名：', '曹宇天'),
    ('学号：', '8002125290'),
    ('专业班级：', '软件工程2510班'),
    ('实验日期：', '2026年5月6日'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.font.size = Pt(14)
    r1.font.bold = True
    r2 = p.add_run(value)
    r2.font.size = Pt(14)

doc.add_page_break()

# ========== 一、实验目标 ==========
doc.add_heading('一、实验目标', level=2)
doc.add_paragraph(
    '目标1. 能够使用Java语言和面向对象程序设计方法，设计和实现满足特定需求的简易学生信息系统，'
    '能根据给定的需求进行分析、设计合理的类，并编码实现，设计的系统包括类与类之间的组合、'
    '继承关系。在设计中体现创新意识。'
)
doc.add_paragraph(
    '目标2. 能够选择与使用恰当的Java开发工具，对简易学生信息系统进行分析、设计，'
    '包括分析和设计UML类图、流程图，调试、优化Java程序等。'
)

# ========== 二、实验内容 ==========
doc.add_heading('二、实验内容', level=2)
doc.add_paragraph(
    '完成一个简易学生信息系统，本系统中学生分为本科生和研究生。本科生有学号、姓名、年龄、'
    '班级、专业、地址、各科成绩等信息，研究生有学号、姓名、年龄、班级、地址、导师、研究方向、'
    '各科成绩等信息，所有地址包含省份、城市、街道、门牌号等信息。系统实现以下功能：'
)

funcs = [
    '增加、修改、删除学生信息',
    '浏览（各类）学生信息（全部/仅本科生/仅研究生）',
    '根据班级、姓名或者学号查询学生信息',
    '根据各科成绩、总成绩、学号排序显示（各类）所有学生信息',
    '分别查看本科生、研究生人数，以及所有学生的总人数',
    'Excel文件（.xlsx）导入与导出',
    'JavaFX图形化界面，支持学生端/老师端双角色入口',
    '老师端登录验证与密码修改功能',
    '数据列表区域拖拽调整高度',
    '智能导入分析器：自动识别表头、字段同义词匹配、容错解析',
]
for i, func in enumerate(funcs, 1):
    doc.add_paragraph(f'{i}. {func}')

doc.add_paragraph()
other_reqs = [
    '（0）本次实验使用AI工具开发（GitHub Copilot、Continue、GPT-5-3-Codex、GPT-5-2、Deepseek v4 pro），尽可能少的手写代码，能精准给AI提清晰需求，能读懂、评审、调优、精简AI生成的代码。',
    '（1）应用面向对象程序设计的思想，根据给定的需求进行分析，设计、实现合理的类。',
    '（2）类间有继承（Student → Undergraduate/Postgraduate）、组合（Student → Address）关系，已绘制类图。',
    '（3）源代码为多文件程序（共12个Java源文件 + 1个module-info.java + 1个CSS样式文件）。',
    '（4）统计学生人数的变量是类中的静态数据成员（Student.totalCount、Undergraduate.count、Postgraduate.count）。',
    '（5）增加学生时检查容量是否已满（isFull()），删除学生时检查仓储是否为空（isEmpty()），边界条件处理完善。',
    '（6）代码规范、美观，易读、易扩展。每个文件、类、方法均有Javadoc注释。',
    '（7）具备完整的图形化菜单功能，包含"关于"项，显示程序名称、版本号、完成时间、姓名、学号、班级等信息。',
]
for req in other_reqs:
    doc.add_paragraph(req)

doc.add_page_break()

# ========== 三、实验环境 ==========
doc.add_heading('三、实验环境', level=2)
env_items = [
    '1. 使用的操作系统及版本：Windows 11 Home 10.0.26200 (25H2)',
    '',
    '2. 使用的开发工具及版本号（包括所用到的AI、插件等）：',
    '   开发工具：IntelliJ IDEA 2025.3.3',
    '   Java SDK：JDK 25.0.2',
    '   JavaFX SDK：javafx-sdk-26',
    '   构建脚本：PowerShell build.ps1',
    '',
    '   插件：GitHub Copilot、Continue',
    '   AI模型：GPT-5-3-Codex、GPT-5-2、Deepseek v4 pro',
    '   开发模式：AI辅助开发，通过精准提示词驱动代码生成，人工评审与调优',
]
for item in env_items:
    doc.add_paragraph(item)

# ========== 四、类图 ==========
doc.add_heading('四、类图（并画出类之间的关系）', level=2)

p = doc.add_paragraph()
run = p.add_run('1. 类图制作工具：')
run.bold = True
p.add_run('IntelliJ IDEA Diagram 插件（UML类图生成）')

p = doc.add_paragraph()
run = p.add_run('2. 系统类设计如下：')
run.bold = True

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '类名'
hdr[1].text = '说明'
hdr[2].text = '与其他类的关系'

class_info = [
    ('Student（抽象类）',
     '学生基类，定义学号、姓名、年龄、班级、地址(Address类型)、\n'
     '成绩Map等通用属性与方法。含静态成员 totalCount。\n'
     '抽象方法：getStudentType()、getTypeSpecificInfo()',
     '父类，被 Undergraduate 和 Postgraduate 继承\n'
     '组合 Address 对象'),
    ('Undergraduate',
     '本科生实体类，继承 Student。\n'
     '新增属性：major（专业）。\n'
     '静态成员 count 统计本科生人数。',
     '继承 Student'),
    ('Postgraduate',
     '研究生实体类，继承 Student。\n'
     '新增属性：supervisor（导师）、researchDirection（研究方向）。\n'
     '静态成员 count 统计研究生人数。',
     '继承 Student'),
    ('Address',
     '地址信息类，封装省份province、城市city、\n'
     '街道street、门牌号houseNumber。\n'
     '重写 equals()/hashCode()/toString()。',
     '被 Student 组合（Student 持有 Address 成员变量）'),
    ('StudentRepository',
     '学生仓储层（数据访问），负责基础数据存储与查找。\n'
     '内部使用 ArrayList<Student>，容量限制2000。\n'
     '处理边界条件：isFull()、isEmpty()。',
     '持有 List<Student> 集合'),
    ('StudentService',
     '学生业务服务层，封装核心业务逻辑。\n'
     '包含：增删改查、排序、统计、边界检查。\n'
     '持有 StudentRepository 引用。',
     '依赖 StudentRepository'),
    ('MenuController',
     '业务控制器，封装界面层调用的业务逻辑。\n'
     '提供 OperationResult/QueryResult/StudentFormData 内部类。\n'
     '持有 StudentService 引用。',
     '依赖 StudentService'),
    ('MainApp',
     'JavaFX 图形界面主类（继承 javafx.application.Application）。\n'
     '实现完整 GUI 交互、角色权限控制、文件导入导出。\n'
     '流程：启动页 -> 双角色入口 -> 主业务界面。',
     '依赖 MenuController、StudentService\n'
     '继承 Application'),
    ('Main',
     '程序入口类，初始化 StudentRepository(2000)、\n'
     'StudentService 并调用 MainApp.launchApp() 启动界面。',
     '调用 MainApp'),
    ('StudentType（枚举）',
     '定义学生类型：UNDERGRADUATE（"本科生"）、POSTGRADUATE（"研究生"）。',
     '被 Student 及其子类使用'),
    ('SortField（枚举）',
     '定义排序字段：STUDENT_ID、TOTAL_SCORE、SUBJECT_SCORE。',
     '被 StudentService.sortStudents() 使用'),
    ('SimpleXlsxReader',
     '简易 XLSX 读取工具，基于 Zip + XML 解析。\n'
     '支持多工作表自动选择、共享字符串表解析。',
     '独立工具类'),
    ('SimpleXlsxWriter',
     '简易 XLSX 写入工具，生成标准 OOXML 格式文件。\n'
     '包含 [Content_Types].xml、workbook.xml、worksheet.xml 等。',
     '独立工具类'),
    ('ImportSheetAnalyzer',
     '智能导入分析器，自动识别表头、字段同义词匹配、\n'
     '成绩列检测、排名列排除、数据容错。\n'
     '使用 AnalysisResult 内部类封装分析结果。',
     '独立工具类，被 MainApp 导入流程调用'),
    ('InputHelper',
     '控制台输入辅助工具，提供类型安全的用户输入读取方法。\n'
     '包含 readNonBlank()、readIntInRange()、readDoubleInRange()、\n'
     'readYesNo()、readAddress()、readScores() 等。',
     '独立工具类（控制台模式备用）'),
]

for cls, desc, rel in class_info:
    row = table.add_row()
    row.cells[0].text = cls
    row.cells[1].text = desc
    row.cells[2].text = rel

p = doc.add_paragraph()
p.add_run('\n3. 类关系图（简化的UML文本表示）：\n').bold = True

uml_text = """\
类之间的关系说明：
  ◇── 组合 Composition（Student 包含 Address）
  ──▷ 继承 Generalization（Undergraduate/Postgraduate 继承 Student）
  ──→ 依赖 Dependency（Service 依赖 Repository，Controller 依赖 Service）

继承层次:
  Student (abstract)
    ├── Undergraduate (+major, +static count)
    └── Postgraduate (+supervisor, +researchDirection, +static count)

组合关系:
  Student ◇── Address (province, city, street, houseNumber)

分层依赖:
  Main → MainApp → MenuController → StudentService → StudentRepository → List<Student>

枚举类型:
  StudentType { UNDERGRADUATE, POSTGRADUATE }
  SortField { STUDENT_ID, TOTAL_SCORE, SUBJECT_SCORE }

工具类:
  SimpleXlsxReader, SimpleXlsxWriter, ImportSheetAnalyzer, InputHelper
"""
p = doc.add_paragraph()
p.add_run(uml_text).font.size = Pt(10)

doc.add_page_break()

# ========== 五、实验步骤及说明 ==========
doc.add_heading('五、实验步骤及说明', level=2)

# 源码列表
p = doc.add_paragraph()
run = p.add_run('本系统源代码文件清单（共15个文件）：')
run.bold = True

files_list = [
    '1. module-info.java —— 模块声明文件，声明 javafx.controls、java.prefs 依赖',
    '2. Main.java —— 程序入口，初始化 StudentRepository(2000) 和 StudentService',
    '3. MainApp.java —— JavaFX 图形界面主类（约3000+行），实现完整 GUI',
    '4. MenuController.java —— 业务控制器（488行），封装界面层业务逻辑',
    '5. StudentService.java —— 业务服务层（285行），核心业务逻辑',
    '6. StudentRepository.java —— 数据仓储层（151行），存储与查找',
    '7. Student.java —— 学生抽象基类（277行），定义通用属性与方法',
    '8. Undergraduate.java —— 本科生实体类（105行），继承 Student',
    '9. Postgraduate.java —— 研究生实体类（130行），继承 Student',
    '10. Address.java —— 地址信息类（159行），封装省份城市街道门牌号',
    '11. StudentType.java —— 学生类型枚举（37行）',
    '12. SortField.java —— 排序字段枚举（15行）',
    '13. SimpleXlsxReader.java —— XLSX读取工具（426行）',
    '14. SimpleXlsxWriter.java —— XLSX写入工具（233行）',
    '15. ImportSheetAnalyzer.java —— 智能导入分析器（800行）',
    '16. InputHelper.java —— 控制台输入辅助（147行）',
    '17. resources/style.css —— JavaFX样式表（436行）',
]
for f in files_list:
    doc.add_paragraph(f)

doc.add_paragraph()

steps = [
    ('步骤1：需求分析与类设计',
     '根据实验需求进行分析，梳理出学生信息系统的核心实体：学生（分为本科生和研究生）、地址。'
     '确定类之间的关系：Student为抽象基类，Undergraduate和Postgraduate继承Student（"is-a"关系）；'
     'Student组合Address对象（"has-a"关系）。同时设计StudentType和SortField枚举类。'
     '设计分层架构：数据仓储层(StudentRepository) → 业务服务层(StudentService) → '
     '控制层(MenuController) → 界面层(MainApp)，各层职责明确、低耦合高内聚。'),

    ('步骤2：搭建项目结构与开发环境',
     '使用IntelliJ IDEA 2025.3.3创建Java项目，配置JavaFX SDK 26与JDK 25.0.2。'
     '设置module-info.java声明模块依赖（requires javafx.controls, requires java.prefs）。'
     '创建源代码目录结构，所有源文件位于studentmanage包下。'
     '编写build.ps1编译运行脚本，支持一键编译与启动（编译-运行分离，自动复制资源文件）。'),

    ('步骤3：实现数据模型层',
     '实现Address类（159行）：封装province、city、street、houseNumber四个属性，'
     '提供完整getter/setter，重写equals()、hashCode()和toString()。'
     '实现Student抽象基类（277行）：定义studentId、name、age、className、address(Address)、'
     'subjectScores(Map<String,Double>)等属性，以及静态成员totalCount（构造时自动递增）。'
     '提供getTotalScore()计算总分、getAverageScore()计算平均分、setSubjectScore()设置单科成绩。'
     '定义抽象方法getStudentType()和getTypeSpecificInfo()供子类多态实现。'
     '实现Undergraduate类（105行）：新增major属性、静态计数器count。'
     '实现Postgraduate类（130行）：新增supervisor、researchDirection属性、静态计数器count。'
     '实现StudentType枚举和SortField枚举。'),

    ('步骤4：实现StudentRepository仓储层',
     '内部使用ArrayList<Student>存储数据，容量参数从构造函数传入（设置2000）。'
     'add()方法：先检查isFull()边界条件（容量满时拒绝新增）。'
     'removeById()方法：先检查isEmpty()边界条件（仓储空时直接返回false），'
     '删除成功后同步维护三级静态计数器。'
     'findById()返回Optional<Student>。getAll()返回列表防御性副本。clearAll()清空并重置计数器。'),

    ('步骤5：实现StudentService业务层',
     '持有StudentRepository引用，封装全部业务逻辑。'
     'addStudent()：先检查学号唯一性（已存在则拒绝）。'
     'updateStudentBasicInfo()：查找后更新共有字段，通过instanceof判断类型分别更新扩展字段。'
     '查询方法：findById()精确匹配、findByName()模糊匹配（toLowerCase+contains）、findByClassName()忽略大小写。'
     'sortStudents()：使用Comparator链式调用，支持按学号/总分/单科成绩排序，升序/降序，学号作为二级排序键。'
     '统计方法：getUndergraduateCount()、getPostgraduateCount()、getTotalCount()。'),

    ('步骤6：实现MenuController控制层',
     '持有StudentService引用，为界面层提供统一调用接口。'
     'addStudent(StudentFormData)：先检查isFull()，根据type构造对应对象，失败时回滚计数器。'
     'updateStudent()：检查存在性和跨类型限制。deleteStudentById()：检查空学号和空仓储。'
     'queryStudents()：支持三种模式。sortStudents()：支持范围和排序方式组合。'
     '设计三个静态内部类：OperationResult、QueryResult、StudentFormData。'),

    ('步骤7：实现JavaFX图形界面（MainApp，约3000+行）',
     '启动页(buildStartupScene)：背景图 + 圆形裁剪头像 + 宋体标题 + 学生端/老师端双按钮。'
     '老师端登录(showTeacherLoginDialog)：管理员账号验证 + 修改密码(showTeacherPasswordChangeDialog)。'
     '主界面：顶部标题栏 + 中间SplitPane(左数据列表+右操作面板) + 底部状态栏。'
     '右侧操作面板使用Accordion折叠块组织：学生信息表单、文件导入、查询排序、统计信息、关于信息。'
     '表单：类型选择 + 基本信息字段 + 地址四字段(省/市联动) + 类型扩展字段 + 动态科目成绩行。'
     '数据列表TableView：动态列生成，成绩低分标红，右键菜单删除，底部拖拽热区调整高度。'
     'CSS样式美化(resources/style.css，436行)：统一字体/颜色/按钮/表格/输入框/状态栏风格。'),

    ('步骤8：实现文件导入导出',
     'SimpleXlsxReader（426行）：基于Zip+XML解析.xlsx，支持sharedStrings表、多工作表自动选择。'
     'SimpleXlsxWriter（233行）：生成标准OOXML格式ZIP文件（含Content_Types、rels、workbook、worksheet）。'
     'ImportSheetAnalyzer（800行）：智能分析——自动定位表头行（扫描40-60行评分）、'
     '字段同义词匹配（学号/考号/准考证号→id）、成绩列检测（数值占比+排除排名/等级列）、'
     '空行/说明行过滤、成绩容错（全角数字/百分号/等级分/缺考→0分）。'
     '导出支持CSV和XLSX两种格式。'),

    ('步骤9：测试与调试',
     '编译验证（build.ps1编译通过，无阻塞错误）。'
     '功能测试：新增/修改/删除学生、按条件查询、多种排序、统计人数、'
     '边界条件（满容量新增、空仓储删除、重复学号拒绝）、导入异常文件。'
     '修复了多个问题：MainApp直接启动空指针、导入分析器误判成绩列（排名列/等级列）、'
     '重复总分列、滚动条缺失与文字漂移、清空表单入口消失、启动参数冲突、容量扩展至2000。'),

    ('步骤10：编写文档',
     '编写实验报告、README.md项目说明、DevNotes.md开发日志（约500行，记录全部关键对话）。'),
]

for title_text, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(13)
    p = doc.add_paragraph()
    p.add_run(desc)
    doc.add_paragraph()

doc.add_page_break()

# ========== 六、实验小结及思考 ==========
doc.add_heading('六、实验小结及思考', level=2)

thoughts = [
    ('1. 面向对象设计能力的提升',
     '通过本次实验深刻理解了面向对象程序设计的核心思想。将Student设计为抽象基类，'
     'Undergraduate和Postgraduate继承并扩展各自特有属性，体现了"is-a"继承关系；'
     'Student组合Address体现了"has-a"组合关系。静态数据成员使人数统计简洁高效——'
     '构造时自动递增计数，删除时同步递减。分层架构（Repository → Service → Controller → View）'
     '使各层职责清晰、低耦合高内聚，例如增加新学生类型只需新增子类和枚举值，业务逻辑无需大改。'),

    ('2. AI辅助开发模式的实践',
     '本次实验使用AI工具开发，学会了如何精准描述需求（用角色+场景+约束的格式写prompt），'
     '如何评审AI生成的代码（检查边界条件、异常处理、逻辑正确性），'
     '如何调优和精简代码（去除冗余、合并重复逻辑）。'
     'AI工具在生成样板代码（getter/setter、CRUD操作）和处理复杂算法（导入分析器、XLSX解析）时效率极高。'
     '但AI代码也有局限性：偶尔产生"幻觉"（不存在的API）、过度设计（不必要的抽象）、忽略边界条件等。'
     '深刻体会到：AI是强大的加速器，但方向盘必须掌握在开发者手中。'),

    ('3. JavaFX图形界面开发经验',
     '掌握了Scene管理、多种Layout布局（BorderPane、SplitPane、GridPane、VBox、HBox、'
     'StackPane、ScrollPane、Accordion）、Control控件（TableView、ComboBox、TextField、'
     'TextArea、Button、PasswordField、CheckBox、TitledPane）、CSS样式、Dialog对话框、'
     '事件处理（鼠标点击/拖拽/释放、右键菜单）、偏好设置存储等。'),

    ('4. 文件处理与容错设计',
     'XLSX读写加深了对OOXML格式的理解（ZIP+XML结构）。ImportSheetAnalyzer的设计是核心挑战'
     '——真实成绩单格式千差万别。通过同义词字典、数值占比检测、排名列排除、工作表评分选择等策略，'
     '实现了高兼容性。容错设计也是重要收获：对各种异常格式尽可能容错而非直接失败，同时记录告警供核查。'),

    ('5. 边界条件与防御性编程',
     '特别注意了边界条件处理：新增时检查isFull()（容量2000满）、删除时检查isEmpty()（无数据）、'
     '所有用户输入经非空+范围校验、学号全局唯一性校验、删除后三级计数器同步维护、'
     '清空操作二次确认弹窗。这些细节在实际系统中至关重要。'),

    ('6. 实验建议',
     '（1）建议增加AI代码评审要求，考察学生的代码审查能力和批判性思维。'
     '（2）可增加数据库持久化扩展（如JDBC连接MySQL），体验从内存到数据库的架构迁移。'
     '（3）可增加团队协作要求（2-3人一组），使用Git分支管理和代码合并。'
     '（4）可增加JUnit单元测试要求。'),
]

for title_text, text in thoughts:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    p = doc.add_paragraph()
    p.add_run(text)
    doc.add_paragraph()

# ========== 七、自评分和自评语 ==========
doc.add_heading('七、自评分和自评语', level=2)

p = doc.add_paragraph()
run = p.add_run('自评分：')
run.font.size = Pt(14)
run.font.bold = True
run = p.add_run('88')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.add_run('\n自评语：').bold = True
eval_text = (
    '本次实验较好地达成了所有实验目标。采用面向对象设计方法，实现了完整的类继承体系'
    '（Student抽象类 → Undergraduate/Postgraduate具体类）和组合关系（Student组合Address），'
    '使用了静态成员完成三级人数统计，合理处理了新增/删除的边界条件。'
    '系统功能完整，包含了要求的所有增删改查、查询、排序、统计功能，'
    '并在此基础上扩展了JavaFX图形化界面、双角色入口（学生端/老师端）、'
    'Excel文件导入导出（含智能导入分析器）、数据列表拖拽调整、右键删除操作等增强功能。'
    '代码风格统一规范，每个文件、类、方法均有完整的Javadoc注释。'
    '全程使用AI工具辅助开发，能精准描述需求并评审优化AI生成的代码。'
    '系统创新点在于：智能导入分析器能自动适配不同格式的成绩单文件（不依赖固定列名），'
    '双角色权限设计为后续扩展预留了空间，数据列表区域拖拽调整高度提升了用户体验。'
    '不足：部分界面细节仍有优化空间，当前仅支持内存存储，未实现数据库持久化。'
)
doc.add_paragraph(eval_text)

p = doc.add_paragraph()
p.add_run('\n各项评分明细：').bold = True

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '评分项（满分100分）'
hdr[1].text = '满分'
hdr[2].text = '自评'
hdr[3].text = '评分说明'

score_items = [
    ('1. 使用AI工具开发，代码无冗余，结构工整易读易扩展', '20', '18',
     '精准使用AI提需求并评审优化，代码工整、无冗余、易扩展'),
    ('2. 采用OOP设计（至少运用继承、组合等机制）', '15', '14',
     'Student抽象基类 + 两个子类继承 + Address组合，类图完整'),
    ('3. 编码风格统一规范，命名严谨合规', '10', '8',
     '统一命名规范，每个文件/类/方法前均有Javadoc注释'),
    ('4. 至少完成实验内容中的功能要求', '15', '13',
     '完成所有要求功能，并有丰富的JavaFX GUI扩展'),
    ('5. 增加了额外功能和性能要求', '10', '8',
     '智能导入分析器、双角色权限、拖拽调整、Excel导入导出等'),
    ('6. 菜单功能与"关于"项', '5', '5',
     '完整的图形化菜单，"关于"信息齐全（名称/版本/日期/姓名/学号/班级）'),
    ('7. 界面友好、输入校验', '5', '4',
     'CSS美化界面、操作提示明确、输入数据有合法性检查'),
    ('8. 注释规范完整', '5', '5',
     '每个源文件、类、方法、关键变量均有规范注释（含作者/功能/时间）'),
    ('9. 多文件项目组织', '5', '5',
     '12个Java源文件 + module-info + CSS，职责分明，目录结构清晰'),
    ('10. 特色和创新', '5', '4',
     '智能导入分析器、双角色权限设计、拖拽交互'),
    ('11. 报告的完整性和逻辑性', '5', '4',
     '报告结构完整、逻辑清晰，含类图、实验步骤、小结、自评分'),
]
for item in score_items:
    row = table.add_row()
    row.cells[0].text = item[0]
    row.cells[1].text = item[1]
    row.cells[2].text = item[2]
    row.cells[3].text = item[3]

# 保存
output_path = 'E:/Projects/StudentManage/面向对象程序设计-期中考试实验报告-已完成.docx'
doc.save(output_path)
print(f'OK: {output_path}')
